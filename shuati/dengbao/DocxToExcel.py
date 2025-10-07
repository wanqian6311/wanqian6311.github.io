import re
from docx import Document
import pandas as pd
import os

def extract_questions_by_delimiters(docx_path):
    """
    使用序号.和A.B.C.D.作为分隔符提取题目
    """
    doc = Document(docx_path)
    questions = []
    current_question = {}
    
    # 分隔符模式
    question_delimiter = re.compile(r'^(\d+)\.')  # 数字. 开头
    option_delimiter = re.compile(r'^([A-D])\.')  # A. B. C. D. 开头
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        # 检查是否是题目分隔符（数字.开头）
        if question_delimiter.match(text):
            # 保存前一个题目
            if current_question and 'question' in current_question:
                questions.append(current_question)
            
            # 开始新题目
            parts = text.split('.', 1)  # 按第一个点分割
            if len(parts) == 2:
                current_question = {
                    'question_number': parts[0].strip(),
                    'question': parts[1].strip(),
                    'options': {}
                }
            else:
                current_question = {
                    'question_number': text.replace('.', '').strip(),
                    'question': '',
                    'options': {}
                }
            continue
        
        # 检查是否是选项分隔符（A. B. C. D.开头）
        if option_delimiter.match(text) and current_question:
            parts = text.split('.', 1)  # 按第一个点分割
            if len(parts) == 2:
                option_letter = parts[0].strip()
                option_text = parts[1].strip()
                current_question['options'][option_letter] = option_text
            continue
        
        # 如果不是分隔符，但有当前题目，则作为题目或选项的延续
        if current_question and 'question' in current_question:
            # 检查当前行是否包含选项内容（可能选项跨行了）
            if text and len(current_question['options']) > 0:
                # 如果已经有选项，可能是最后一个选项的延续
                last_option = max(current_question['options'].keys()) if current_question['options'] else None
                if last_option:
                    current_question['options'][last_option] += ' ' + text
            else:
                # 否则作为题目的延续
                current_question['question'] += ' ' + text
    
    # 添加最后一个题目
    if current_question and 'question' in current_question:
        questions.append(current_question)
    
    return questions

def extract_questions_advanced(docx_path):
    """
    更强大的提取方法，处理各种复杂情况
    """
    doc = Document(docx_path)
    full_text = []
    
    # 获取所有非空文本
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            full_text.append(text)
    
    questions = []
    
    # 使用正则表达式匹配题目和选项
    # 匹配模式：数字. 题目内容 A. 选项A B. 选项B C. 选项C D. 选项D
    pattern = re.compile(r'(\d+)\.\s*(.*?)\s*(A\.\s*.*?)\s*(B\.\s*.*?)\s*(C\.\s*.*?)\s*(D\.\s*.*?)(?=\s*\d+\.|$)', re.DOTALL)
    
    combined_text = ' '.join(full_text)
    matches = pattern.findall(combined_text)
    
    for match in matches:
        question_num = match[0]
        question_text = match[1].strip()
        
        # 提取选项
        options = {}
        for i in range(2, 6):  # A, B, C, D 选项
            option_text = match[i]
            # 去除选项标识符并清理
            option_clean = re.sub(r'^[A-D]\.\s*', '', option_text).strip()
            option_letter = chr(65 + i - 2)  # A, B, C, D
            options[option_letter] = option_clean
        
        questions.append({
            'question_number': question_num,
            'question': question_text,
            'options': options
        })
    
    # 如果上面的方法没找到，尝试逐段分析
    if not questions:
        questions = extract_questions_by_delimiters(docx_path)
    
    return questions

def process_questions_to_dataframe(questions):
    """
    将提取的题目转换为DataFrame
    """
    data = []
    
    for q in questions:
        row = {
            '题号': q['question_number'],
            '题目': q['question'],
            '选项A': q['options'].get('A', ''),
            '选项B': q['options'].get('B', ''),
            '选项C': q['options'].get('C', ''),
            '选项D': q['options'].get('D', '')
        }
        data.append(row)
    
    df = pd.DataFrame(data)
    return df

def word_to_excel(input_path, output_path=None):
    """
    主函数：将Word转换为Excel
    """
    if not os.path.exists(input_path):
        print(f"错误：文件 '{input_path}' 不存在")
        return False
    
    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}.xlsx"
    
    try:
        print("正在从Word文档中提取题目...")
        
        # 首先尝试高级提取
        questions = extract_questions_advanced(input_path)
        
        # 如果没找到，使用基础分隔符方法
        if not questions:
            print("高级提取未找到题目，尝试基础分隔符提取...")
            questions = extract_questions_by_delimiters(input_path)
        
        if not questions:
            print("未找到任何题目，请检查文档格式")
            print("期望格式: 1. 题目内容 A. 选项1 B. 选项2 C. 选项3 D. 选项4")
            return False
        
        print(f"成功提取 {len(questions)} 道题目")
        
        # 显示前3题作为示例
        print("\n前3题提取结果示例:")
        for i, q in enumerate(questions[:3]):
            print(f"第{q['question_number']}题: {q['question']}")
            for option in ['A', 'B', 'C', 'D']:
                if option in q['options']:
                    print(f"  {option}. {q['options'][option]}")
            print()
        
        # 转换为DataFrame并保存
        df = process_questions_to_dataframe(questions)
        df.to_excel(output_path, index=False, engine='openpyxl')
        
        print(f"成功将题目保存到: {output_path}")
        return True
        
    except Exception as e:
        print(f"转换过程中出现错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def analyze_document_format(docx_path):
    """
    分析文档格式，帮助调试
    """
    doc = Document(docx_path)
    print("=== 文档格式分析 ===")
    
    question_count = 0
    option_count = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
            
        print(f"段落 {i}: {text[:50]}...")  # 显示前50个字符
        
        # 统计题目和选项
        if re.match(r'^\d+\.', text):
            question_count += 1
            print(f"  → 识别为题目")
        elif re.match(r'^[A-D]\.', text):
            option_count += 1
            print(f"  → 识别为选项")
    
    print(f"\n统计: 发现 {question_count} 个题目起始, {option_count} 个选项起始")

# 使用示例
if __name__ == "__main__":
    # 请修改为您的实际文件路径
    input_file = "网络安全等级测评师能力评估（初级）样卷3.docx"
    output_file = "网络安全等级测评师能力评估（初级）样卷3.xlsx"
    
    # 先分析文档格式
    analyze_document_format(input_file)
    
    # 执行转换
    success = word_to_excel(input_file, output_file)
    
    if success:
        print("\n🎉 转换完成！")
    else:
        print("\n❌ 转换失败！")
