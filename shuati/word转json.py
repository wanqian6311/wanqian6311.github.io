from docx import Document
import json

def extract_questions(doc_path):
    doc = Document(doc_path)
    questions = []
    question = None

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        if text.startswith("题目"):
            if question:  # 如果之前有题目，保存之前的题目
                questions.append(question)
            question = {
                "id": str(len(questions) + 1),  # 根据题目的顺序生成 id
                "title": "",  # 初始化标题
                "option": [],
                "answer": "",
                "analysis": ""  
            }
        elif question is not None and question["title"] == "":
            # 如果当前是题目且标题尚未设置，则将下一行设置为标题
            question["title"] = text
        elif text.startswith("A．") or text.startswith("B．") or text.startswith("C．") or text.startswith("D．"):
            question["option"].append(text)
        elif text.startswith("答案"):
            # 如果当前是答案，则将下一行设置为答案
            if i + 1 < len(doc.paragraphs):
                question["answer"] = doc.paragraphs[i + 1].text.strip()
        elif text.startswith("详解"):
            # 如果当前是详解，则将下一行设置为详解
            if i + 1 < len(doc.paragraphs):
                question["analysis"] = doc.paragraphs[i + 1].text.strip()
                questions.append(question)
                question = None  # 重置题目

    return questions

def save_to_json(data, json_path):
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

if __name__ == "__main__":
    doc_path = './多选.docx'
    questions = extract_questions(doc_path)
    save_to_json(questions, 'questions3.json')
    print("Successfully converted Word document to JSON: questions3.json")
